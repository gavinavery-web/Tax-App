"""Text extraction for the Evidence Vault.

Supports PDF (text-layer + OCR fallback), images, DOCX, XLSX, CSV, TXT.
Returns the extracted plain text. Truncated to a 20,000-char ceiling for the AI
call separately (see ai_classifier.truncate_for_ai).
"""
from __future__ import annotations

import io
import logging
import os
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

MAX_AI_CHARS = 20000


def _ocr_image_bytes(img_bytes: bytes) -> str:
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(io.BytesIO(img_bytes))
        return pytesseract.image_to_string(img)
    except Exception as e:
        logger.warning(f"OCR failed: {e}")
        return ""


def _extract_pdf(content: bytes) -> str:
    """Extract text from PDF. Use pdfplumber first; OCR pages with little text."""
    import pdfplumber

    out_pages: list[str] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for i, page in enumerate(pdf.pages):
            txt = page.extract_text() or ""
            if len(txt.strip()) < 50:
                # likely scanned — OCR fallback via pdf2image
                try:
                    from pdf2image import convert_from_bytes
                    images = convert_from_bytes(content, first_page=i + 1, last_page=i + 1, dpi=200)
                    if images:
                        buf = io.BytesIO()
                        images[0].save(buf, format="PNG")
                        txt = _ocr_image_bytes(buf.getvalue())
                except Exception as e:
                    logger.warning(f"PDF OCR page {i+1} failed: {e}")
            out_pages.append(txt or "")
    return "\n\n".join(out_pages)


def _extract_docx(content: bytes) -> str:
    from docx import Document as DocxDocument
    doc = DocxDocument(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(parts)


def _extract_xlsx(content: bytes) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(content), data_only=True, read_only=True)
    parts: list[str] = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        parts.append(f"## Sheet: {sheet}")
        for row in ws.iter_rows(values_only=True):
            line = "\t".join("" if v is None else str(v) for v in row)
            if line.strip():
                parts.append(line)
    return "\n".join(parts)


def _extract_csv(content: bytes) -> str:
    try:
        return content.decode("utf-8", errors="replace")
    except Exception:
        return content.decode("latin-1", errors="replace")


def _extract_image(content: bytes) -> str:
    return _ocr_image_bytes(content)


def extract_text(content: bytes, filename: str, mime: Optional[str] = None) -> str:
    ext = (Path(filename).suffix or "").lower().lstrip(".")
    mime = (mime or "").lower()
    try:
        if ext == "pdf" or "pdf" in mime:
            return _extract_pdf(content)
        if ext in ("docx", "doc"):
            return _extract_docx(content)
        if ext in ("xlsx", "xls"):
            return _extract_xlsx(content)
        if ext == "csv" or "csv" in mime:
            return _extract_csv(content)
        if ext in ("png", "jpg", "jpeg", "webp", "heic", "heif", "tiff", "bmp") or mime.startswith("image/"):
            return _extract_image(content)
        if ext == "txt" or mime.startswith("text/"):
            return _extract_csv(content)
    except Exception as e:
        logger.exception(f"Extraction failed for {filename}: {e}")
        return ""
    return ""


def truncate_for_ai(text: str, limit: int = MAX_AI_CHARS) -> str:
    if not text:
        return ""
    if len(text) <= limit:
        return text
    head = limit * 12 // 20  # 12/20 ≈ 60%
    tail = limit - head - 30
    return text[:head] + "\n\n[... truncated for AI ...]\n\n" + text[-tail:]
